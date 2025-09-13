"""
Document Generation Agent for HR Advisor
Handles automated generation of HR documents including contracts, policies, forms, and templates
"""

import asyncio
import json
import uuid
from datetime import datetime, timedelta
from typing import Dict, List, Any, Optional
from dataclasses import dataclass
from enum import Enum
import os
from jinja2 import Template, Environment, FileSystemLoader
import pdfkit
import docx
from docx import Document
from docx.shared import Inches
import markdown

class DocumentType(Enum):
    EMPLOYMENT_CONTRACT = "employment_contract"
    OFFER_LETTER = "offer_letter"
    TERMINATION_LETTER = "termination_letter"
    POLICY_DOCUMENT = "policy_document"
    PERFORMANCE_REVIEW = "performance_review"
    WARNING_LETTER = "warning_letter"
    REFERENCE_LETTER = "reference_letter"
    JOB_DESCRIPTION = "job_description"
    HANDBOOK = "employee_handbook"
    TRAINING_MATERIAL = "training_material"
    COMPLIANCE_REPORT = "compliance_report"
    ONBOARDING_CHECKLIST = "onboarding_checklist"

class DocumentFormat(Enum):
    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"
    MARKDOWN = "markdown"
    TXT = "txt"

@dataclass
class DocumentTemplate:
    template_id: str
    name: str
    document_type: DocumentType
    country: str
    language: str
    template_content: str
    required_fields: List[str]
    optional_fields: List[str]
    legal_compliance: List[str]
    version: str
    created_date: datetime
    updated_date: datetime

@dataclass
class GeneratedDocument:
    document_id: str
    template_id: str
    document_type: DocumentType
    title: str
    content: str
    format: DocumentFormat
    employee_id: Optional[str]
    generated_by: str
    generated_date: datetime
    file_path: Optional[str]
    metadata: Dict[str, Any]

class DocumentGenerationAgent:
    def __init__(self):
        self.templates: Dict[str, DocumentTemplate] = {}
        self.generated_documents: Dict[str, GeneratedDocument] = {}
        self._initialize_templates()
        
    def _initialize_templates(self):
        """Initialize predefined document templates"""
        
        # Employment Contract Template (US)
        employment_contract_us = DocumentTemplate(
            template_id="emp_contract_us_001",
            name="US Employment Contract - Full Time",
            document_type=DocumentType.EMPLOYMENT_CONTRACT,
            country="US",
            language="en",
            template_content="""
# EMPLOYMENT AGREEMENT

**Between:** {{ company_name }}  
**And:** {{ employee_name }}

## 1. POSITION AND DUTIES
Employee is hired as {{ job_title }} in the {{ department }} department, reporting to {{ manager_name }}.

**Start Date:** {{ start_date }}  
**Employment Type:** {{ employment_type }}  
**Work Location:** {{ work_location }}

## 2. COMPENSATION
**Base Salary:** ${{ annual_salary }} per year, paid {{ pay_frequency }}  
**Benefits:** As outlined in the Employee Handbook

## 3. WORKING HOURS
Standard working hours are {{ working_hours }} per week, {{ work_schedule }}.

## 4. PROBATIONARY PERIOD
Employee will serve a probationary period of {{ probation_period }} months from the start date.

## 5. CONFIDENTIALITY
Employee agrees to maintain confidentiality of all proprietary information and trade secrets.

## 6. TERMINATION
Either party may terminate this agreement with {{ notice_period }} days written notice.

## 7. GOVERNING LAW
This agreement is governed by the laws of {{ state }}, United States.

**Company Representative:**  
Name: {{ company_representative }}  
Title: {{ representative_title }}  
Date: {{ contract_date }}

**Employee:**  
Name: {{ employee_name }}  
Date: ________________

**Signature:** ________________
            """,
            required_fields=[
                "company_name", "employee_name", "job_title", "department", 
                "manager_name", "start_date", "employment_type", "work_location",
                "annual_salary", "pay_frequency", "working_hours", "work_schedule",
                "probation_period", "notice_period", "state", "company_representative",
                "representative_title", "contract_date"
            ],
            optional_fields=["bonus_structure", "stock_options", "additional_benefits"],
            legal_compliance=["FLSA", "Equal Employment Opportunity", "At-Will Employment"],
            version="1.0",
            created_date=datetime.now(),
            updated_date=datetime.now()
        )
        
        # Offer Letter Template
        offer_letter = DocumentTemplate(
            template_id="offer_letter_001",
            name="Job Offer Letter",
            document_type=DocumentType.OFFER_LETTER,
            country="US",
            language="en",
            template_content="""
# JOB OFFER LETTER

{{ company_letterhead }}

{{ date }}

{{ candidate_name }}  
{{ candidate_address }}

Dear {{ candidate_name }},

We are pleased to offer you the position of **{{ job_title }}** with {{ company_name }}.

## Position Details:
- **Department:** {{ department }}
- **Reporting Manager:** {{ manager_name }}
- **Start Date:** {{ start_date }}
- **Work Location:** {{ work_location }}
- **Employment Type:** {{ employment_type }}

## Compensation Package:
- **Annual Salary:** ${{ annual_salary }}
- **Benefits:** {{ benefits_summary }}
- **Vacation:** {{ vacation_days }} days per year

## Next Steps:
Please confirm your acceptance by {{ response_deadline }} by signing and returning this letter.

We look forward to welcoming you to our team!

Sincerely,

{{ hiring_manager_name }}  
{{ hiring_manager_title }}  
{{ company_name }}

---

**ACCEPTANCE:**

I accept the terms of this offer:

Signature: ________________  Date: ________________

{{ candidate_name }}
            """,
            required_fields=[
                "company_name", "candidate_name", "candidate_address", "date",
                "job_title", "department", "manager_name", "start_date",
                "work_location", "employment_type", "annual_salary",
                "benefits_summary", "vacation_days", "response_deadline",
                "hiring_manager_name", "hiring_manager_title"
            ],
            optional_fields=["company_letterhead", "bonus_details", "relocation_assistance"],
            legal_compliance=["Equal Employment Opportunity"],
            version="1.0",
            created_date=datetime.now(),
            updated_date=datetime.now()
        )
        
        # Performance Review Template
        performance_review = DocumentTemplate(
            template_id="perf_review_001",
            name="Annual Performance Review",
            document_type=DocumentType.PERFORMANCE_REVIEW,
            country="US",
            language="en",
            template_content="""
# PERFORMANCE REVIEW

**Employee:** {{ employee_name }}  
**Position:** {{ job_title }}  
**Department:** {{ department }}  
**Review Period:** {{ review_period }}  
**Review Date:** {{ review_date }}  
**Reviewer:** {{ reviewer_name }}

## 1. OVERALL PERFORMANCE RATING
**Rating:** {{ overall_rating }}/5

## 2. KEY ACHIEVEMENTS
{{ achievements }}

## 3. AREAS OF STRENGTH
{{ strengths }}

## 4. AREAS FOR IMPROVEMENT
{{ improvement_areas }}

## 5. GOALS FOR NEXT PERIOD
{{ future_goals }}

## 6. DEVELOPMENT PLAN
{{ development_plan }}

## 7. MANAGER COMMENTS
{{ manager_comments }}

## 8. EMPLOYEE COMMENTS
{{ employee_comments }}

**Manager Signature:** ________________  Date: ________________

**Employee Signature:** ________________  Date: ________________
            """,
            required_fields=[
                "employee_name", "job_title", "department", "review_period",
                "review_date", "reviewer_name", "overall_rating"
            ],
            optional_fields=[
                "achievements", "strengths", "improvement_areas", 
                "future_goals", "development_plan", "manager_comments", "employee_comments"
            ],
            legal_compliance=["Performance Documentation"],
            version="1.0",
            created_date=datetime.now(),
            updated_date=datetime.now()
        )
        
        # Policy Document Template
        policy_template = DocumentTemplate(
            template_id="policy_001",
            name="HR Policy Template",
            document_type=DocumentType.POLICY_DOCUMENT,
            country="US",
            language="en",
            template_content="""
# {{ policy_title }}

**Policy Number:** {{ policy_number }}  
**Effective Date:** {{ effective_date }}  
**Last Reviewed:** {{ review_date }}  
**Next Review:** {{ next_review_date }}

## 1. PURPOSE
{{ policy_purpose }}

## 2. SCOPE
This policy applies to {{ policy_scope }}.

## 3. POLICY STATEMENT
{{ policy_statement }}

## 4. PROCEDURES
{{ procedures }}

## 5. RESPONSIBILITIES
{{ responsibilities }}

## 6. COMPLIANCE
{{ compliance_requirements }}

## 7. VIOLATIONS
{{ violation_consequences }}

## 8. RELATED POLICIES
{{ related_policies }}

**Approved by:** {{ approver_name }}  
**Title:** {{ approver_title }}  
**Date:** {{ approval_date }}
            """,
            required_fields=[
                "policy_title", "policy_number", "effective_date", "review_date",
                "next_review_date", "policy_purpose", "policy_scope",
                "policy_statement", "approver_name", "approver_title", "approval_date"
            ],
            optional_fields=[
                "procedures", "responsibilities", "compliance_requirements",
                "violation_consequences", "related_policies"
            ],
            legal_compliance=["Policy Documentation Standards"],
            version="1.0",
            created_date=datetime.now(),
            updated_date=datetime.now()
        )
        
        # Store templates
        self.templates[employment_contract_us.template_id] = employment_contract_us
        self.templates[offer_letter.template_id] = offer_letter
        self.templates[performance_review.template_id] = performance_review
        self.templates[policy_template.template_id] = policy_template
    
    async def generate_document(self, template_id: str, data: Dict[str, Any], 
                              format: DocumentFormat = DocumentFormat.PDF,
                              generated_by: str = "System") -> str:
        """Generate a document from template"""
        
        if template_id not in self.templates:
            raise ValueError(f"Template {template_id} not found")
        
        template = self.templates[template_id]
        
        # Validate required fields
        missing_fields = [field for field in template.required_fields if field not in data]
        if missing_fields:
            raise ValueError(f"Missing required fields: {missing_fields}")
        
        # Generate document content
        jinja_template = Template(template.template_content)
        content = jinja_template.render(**data)
        
        # Create document record
        document_id = str(uuid.uuid4())
        document = GeneratedDocument(
            document_id=document_id,
            template_id=template_id,
            document_type=template.document_type,
            title=data.get('document_title', f"{template.name} - {data.get('employee_name', 'Document')}"),
            content=content,
            format=format,
            employee_id=data.get('employee_id'),
            generated_by=generated_by,
            generated_date=datetime.now(),
            file_path=None,
            metadata=data
        )
        
        # Generate file if needed
        if format != DocumentFormat.HTML:
            file_path = await self._generate_file(document, format)
            document.file_path = file_path
        
        self.generated_documents[document_id] = document
        return document_id
    
    async def _generate_file(self, document: GeneratedDocument, format: DocumentFormat) -> str:
        """Generate physical file from document content"""
        
        # Create documents directory if it doesn't exist
        docs_dir = "/tmp/hr_documents"
        os.makedirs(docs_dir, exist_ok=True)
        
        filename = f"{document.document_id}_{document.title.replace(' ', '_')}"
        
        if format == DocumentFormat.PDF:
            file_path = f"{docs_dir}/{filename}.pdf"
            # Convert markdown to HTML first, then to PDF
            html_content = markdown.markdown(document.content)
            pdfkit.from_string(html_content, file_path)
            
        elif format == DocumentFormat.DOCX:
            file_path = f"{docs_dir}/{filename}.docx"
            doc = Document()
            
            # Simple conversion from markdown to docx
            lines = document.content.split('\n')
            for line in lines:
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('**') and line.endswith('**'):
                    p = doc.add_paragraph()
                    p.add_run(line[2:-2]).bold = True
                elif line.strip():
                    doc.add_paragraph(line)
            
            doc.save(file_path)
            
        elif format == DocumentFormat.MARKDOWN:
            file_path = f"{docs_dir}/{filename}.md"
            with open(file_path, 'w') as f:
                f.write(document.content)
                
        elif format == DocumentFormat.TXT:
            file_path = f"{docs_dir}/{filename}.txt"
            # Convert markdown to plain text
            plain_text = document.content.replace('#', '').replace('**', '')
            with open(file_path, 'w') as f:
                f.write(plain_text)
        
        return file_path
    
    def get_templates(self, document_type: Optional[DocumentType] = None, 
                     country: Optional[str] = None) -> List[Dict[str, Any]]:
        """Get available templates"""
        templates = []
        
        for template in self.templates.values():
            if document_type and template.document_type != document_type:
                continue
            if country and template.country != country:
                continue
                
            templates.append({
                'template_id': template.template_id,
                'name': template.name,
                'document_type': template.document_type.value,
                'country': template.country,
                'language': template.language,
                'required_fields': template.required_fields,
                'optional_fields': template.optional_fields,
                'version': template.version
            })
        
        return templates
    
    def get_document(self, document_id: str) -> Optional[Dict[str, Any]]:
        """Get generated document details"""
        if document_id not in self.generated_documents:
            return None
        
        document = self.generated_documents[document_id]
        return {
            'document_id': document.document_id,
            'template_id': document.template_id,
            'document_type': document.document_type.value,
            'title': document.title,
            'content': document.content,
            'format': document.format.value,
            'employee_id': document.employee_id,
            'generated_by': document.generated_by,
            'generated_date': document.generated_date.isoformat(),
            'file_path': document.file_path,
            'metadata': document.metadata
        }
    
    def get_employee_documents(self, employee_id: str) -> List[Dict[str, Any]]:
        """Get all documents for an employee"""
        employee_docs = []
        
        for document in self.generated_documents.values():
            if document.employee_id == employee_id:
                employee_docs.append({
                    'document_id': document.document_id,
                    'title': document.title,
                    'document_type': document.document_type.value,
                    'generated_date': document.generated_date.isoformat(),
                    'format': document.format.value,
                    'file_path': document.file_path
                })
        
        return employee_docs
    
    async def create_custom_template(self, name: str, document_type: DocumentType,
                                   country: str, template_content: str,
                                   required_fields: List[str],
                                   optional_fields: List[str] = None,
                                   legal_compliance: List[str] = None) -> str:
        """Create a custom document template"""
        
        template_id = str(uuid.uuid4())
        template = DocumentTemplate(
            template_id=template_id,
            name=name,
            document_type=document_type,
            country=country,
            language="en",
            template_content=template_content,
            required_fields=required_fields,
            optional_fields=optional_fields or [],
            legal_compliance=legal_compliance or [],
            version="1.0",
            created_date=datetime.now(),
            updated_date=datetime.now()
        )
        
        self.templates[template_id] = template
        return template_id
    
    async def generate_bulk_documents(self, template_id: str, 
                                    employee_data_list: List[Dict[str, Any]],
                                    format: DocumentFormat = DocumentFormat.PDF) -> List[str]:
        """Generate documents for multiple employees"""
        document_ids = []
        
        for employee_data in employee_data_list:
            try:
                document_id = await self.generate_document(
                    template_id=template_id,
                    data=employee_data,
                    format=format,
                    generated_by="Bulk Generation"
                )
                document_ids.append(document_id)
            except Exception as e:
                print(f"Error generating document for employee {employee_data.get('employee_name', 'Unknown')}: {e}")
                continue
        
        return document_ids
    
    def get_compliance_requirements(self, country: str, document_type: DocumentType) -> List[str]:
        """Get compliance requirements for document type and country"""
        compliance_map = {
            ("US", DocumentType.EMPLOYMENT_CONTRACT): [
                "FLSA Compliance", "Equal Employment Opportunity", 
                "At-Will Employment Clause", "State-Specific Requirements"
            ],
            ("UK", DocumentType.EMPLOYMENT_CONTRACT): [
                "Employment Rights Act", "GDPR Compliance",
                "Working Time Regulations", "Statutory Rights"
            ],
            ("SG", DocumentType.EMPLOYMENT_CONTRACT): [
                "Employment Act", "MOM Regulations",
                "CPF Requirements", "Work Pass Conditions"
            ]
        }
        
        return compliance_map.get((country, document_type), [])
    
    async def validate_document_compliance(self, document_id: str) -> Dict[str, Any]:
        """Validate document against compliance requirements"""
        if document_id not in self.generated_documents:
            return {"error": "Document not found"}
        
        document = self.generated_documents[document_id]
        template = self.templates[document.template_id]
        
        compliance_issues = []
        
        # Check required fields
        for field in template.required_fields:
            if field not in document.metadata or not document.metadata[field]:
                compliance_issues.append(f"Missing required field: {field}")
        
        # Check legal compliance requirements
        country_requirements = self.get_compliance_requirements(
            template.country, template.document_type
        )
        
        for requirement in country_requirements:
            if requirement not in template.legal_compliance:
                compliance_issues.append(f"Missing compliance requirement: {requirement}")
        
        return {
            "document_id": document_id,
            "compliant": len(compliance_issues) == 0,
            "issues": compliance_issues,
            "requirements_met": template.legal_compliance,
            "country_requirements": country_requirements
        }

# Global instance
document_agent = DocumentGenerationAgent()

